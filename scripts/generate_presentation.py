"""Génère un document Word de présentation du projet TenderAI BF."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

LOGO_PATH     = "/media/abdazz/New Volume/YULCOM/Logos/Logo-Yulcom-Technologies.png"

PRIMARY_COLOR = RGBColor(0x00, 0x4C, 0x97)   # bleu YULCOM principal
SECONDARY_COLOR = RGBColor(0x00, 0x7A, 0xC3) # bleu YULCOM secondaire
LIGHT_GRAY    = RGBColor(0xF2, 0xF2, 0xF2)
DARK_GRAY     = RGBColor(0x40, 0x40, 0x40)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

BODY_SIZE     = Pt(12)
LINE_SPACING  = WD_LINE_SPACING.ONE_POINT_FIVE


# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_border_bottom(paragraph, color: str = "004C97", size: int = 12) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def apply_spacing(p) -> None:
    p.paragraph_format.line_spacing_rule = LINE_SPACING
    p.paragraph_format.space_after = Pt(6)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = LINE_SPACING
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size      = Pt(16)
        run.font.color.rgb = PRIMARY_COLOR
        add_border_bottom(p)
    elif level == 2:
        run.font.size      = Pt(13)
        run.font.color.rgb = SECONDARY_COLOR
    else:
        run.font.size      = Pt(12)
        run.font.color.rgb = DARK_GRAY


def bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    apply_spacing(p)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = BODY_SIZE
        r.font.color.rgb = PRIMARY_COLOR
        rb = p.add_run(text)
        rb.font.size = BODY_SIZE
    else:
        rb = p.add_run(text)
        rb.font.size = BODY_SIZE


def normal(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    apply_spacing(p)
    for run in p.runs:
        run.font.size = BODY_SIZE


def colored_table_row(table, label: str, value: str, bg: bool = False) -> None:
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    r0 = row.cells[0].paragraphs[0].runs[0]
    r0.bold = True
    r0.font.size = BODY_SIZE
    r0.font.color.rgb = PRIMARY_COLOR
    row.cells[1].paragraphs[0].runs[0].font.size = BODY_SIZE
    if bg:
        set_cell_bg(row.cells[0], "F2F2F2")
        set_cell_bg(row.cells[1], "F2F2F2")


# ── header avec logo ──────────────────────────────────────────────────────────

def add_header(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(LOGO_PATH, height=Inches(0.55))


# ── footer ────────────────────────────────────────────────────────────────────

def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("YULCOM Technologies Burkina Faso  |  TenderAI BF")
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_GRAY


# ── cover page ────────────────────────────────────────────────────────────────

def build_cover(doc: Document) -> None:
    # barre bleue de titre
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run("━" * 80)
    run.font.color.rgb = PRIMARY_COLOR
    run.font.size = Pt(8)

    doc.add_paragraph()

    # label
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("DOCUMENT DE PRESENTATION")
    r.font.size = Pt(12)
    r.font.color.rgb = SECONDARY_COLOR
    r.bold = True
    r.font.all_caps = True

    # titre principal
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(16)
    r = title.add_run("TenderAI BF")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = PRIMARY_COLOR

    # sous-titre
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(6)
    r = sub.add_run(
        "Systeme Multi-Agents de Veille Automatique\n"
        "des Appels d'Offres, Burkina Faso"
    )
    r.font.size = Pt(15)
    r.font.color.rgb = DARK_GRAY

    for _ in range(3):
        doc.add_paragraph()

    # tableau de metadonnees
    meta = doc.add_table(rows=0, cols=2)
    meta.style = "Table Grid"
    meta.columns[0].width = Inches(2.2)
    meta.columns[1].width = Inches(3.8)

    rows_data = [
        ("Proprietaire", "YULCOM Technologies Burkina Faso"),
        ("Auteur",       "Abdoul-Aziz ZOROM, CTO"),
        ("Version",      "1.0, Phase 1"),
        ("Date",         datetime.date.today().strftime("%d %B %Y")),
    ]
    for i, (k, v) in enumerate(rows_data):
        colored_table_row(meta, k, v, bg=(i % 2 == 0))

    doc.add_paragraph()

    # barre basse
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    run2 = p2.add_run("━" * 80)
    run2.font.color.rgb = PRIMARY_COLOR
    run2.font.size = Pt(8)

    doc.add_page_break()


# ── section 1 : présentation ──────────────────────────────────────────────────

def build_presentation(doc: Document) -> None:
    heading(doc, "1. Presentation du Projet")

    normal(doc,
        "TenderAI BF est un systeme multi-agents de veille automatique des appels d'offres "
        "(RFP/AO) dans les domaines IT et ingenierie au Burkina Faso. Developpe par YULCOM "
        "Technologies, il scrute quotidiennement les portails publics de la commande publique "
        "burkinabe, identifie les opportunites pertinentes grace a l'intelligence artificielle, "
        "elimine les doublons, genere un rapport Word structure et le distribue automatiquement "
        "par email aux equipes commerciales et Bid Managers."
    )

    heading(doc, "1.1 Contexte et Problematique", 2)
    normal(doc,
        "Le marche de la commande publique burkinabe est fragmente sur de nombreux portails "
        "(ARCOP, DGCMEF, ministeres, bailleurs internationaux). Suivre manuellement ces sources "
        "est chronophage, source d'omissions et peu scalable. TenderAI BF automatise entierement "
        "cette veille, 24h/24 et 7j/7, permettant aux equipes de se concentrer sur la redaction "
        "des offres plutot que sur la collecte d'informations."
    )

    heading(doc, "1.2 Objectifs Principaux", 2)
    bullets = [
        ("Decouverte quotidienne : ",
         "detecter toutes les nouvelles opportunites AO/RFP a partir de sources publiques fiables."),
        ("Classification automatique : ",
         "filtrer les avis pertinents IT/ingenierie avec un rappel >= 90 % et un taux de faux positifs <= 10 %."),
        ("Extraction structuree : ",
         "normaliser les champs cles (titre, entite, dates, budget, contacts) quelle que soit la source."),
        ("Rapport professionnel : ",
         "generer chaque jour a 07h30 un fichier .docx brande YULCOM et l'envoyer par email."),
        ("Supervision en temps reel : ",
         "offrir un tableau de bord Gradio pour relances manuelles, configuration et suivi."),
    ]
    for prefix, text in bullets:
        bullet(doc, text, prefix)

    heading(doc, "1.3 Perimetre : Phase 1", 2)
    normal(doc, "Inclus :")
    for item in [
        "Crawling des portails publics burkinabe (ARCOP, DGCMEF, bailleurs, etc.)",
        "Extraction HTML et PDF avec OCR pour les documents image",
        "Classification hybride regles + LLM",
        "Deduplication inter-sources avec tracabilite de provenance",
        "Generation du rapport .docx et distribution SMTP",
        "Interface d'administration Gradio (sources, historique, destinataires)",
        "Deploiement Docker Compose sur VPS YULCOM",
    ]:
        bullet(doc, item)

    normal(doc, "Hors perimetre (Phase 1) :")
    for item in [
        "Achat de dossiers d'appels d'offres et aide a la redaction des offres",
        "Portails necessitant une authentification non fournie",
        "Support multilingue complet au-dela du francais/anglais",
    ]:
        bullet(doc, item)

    doc.add_page_break()


# ── section 2 : fonctionnalités ───────────────────────────────────────────────

def build_fonctionnalites(doc: Document) -> None:
    heading(doc, "2. Fonctionnalites Cles")

    features = [
        ("Veille multi-sources",
         "Surveillance continue de portails nationaux (ARCOP, DGCMEF), sites ministeriels, "
         "agregateurs specialises et bailleurs internationaux (ONU, UE). Chaque source dispose "
         "de selecteurs CSS/XPath configurables, de limites de debit et d'un suivi last-seen."),
        ("Extraction intelligente",
         "Parsing HTML via selectolax/BeautifulSoup et extraction PDF via pdfminer-six. "
         "Pour les PDF image-only (scan), Docling assure l'OCR en langue francaise. "
         "Les champs extraits (titre, reference, entite, dates, budget, contacts) sont normalises "
         "au format ISO/XOF/region."),
        ("Classification IA hybride",
         "Premiere passe par regles a base de mots-cles metier (logiciel, reseau, cloud, "
         "cybersecurite, data, IA, ERP, SIG...). Deuxieme passe par un LLM (Groq/llama-3.3 par "
         "defaut, OpenAI ou Ollama en option) pour les cas ambigus. Score de confiance ajustable "
         "via le dashboard."),
        ("Deduplication multi-strategies",
         "Empreintes de contenu (hash SHA-256 + similarite textuelle rapidfuzz). Cinq modes "
         "configurables : hash_only, similarity_only, hash_similarity, llm_only, hybrid. "
         "Conservation de la provenance multi-URLs pour l'audit."),
        ("Rapport Word professionnel",
         "Generation automatique d'un .docx avec page de garde, resume executif (KPIs), "
         "table des matieres, fiches AO detaillees et annexes (couverture sources, erreurs). "
         "Nommage : RFP_Watch_BF_YYYY-MM-DD-HH-MM.docx."),
        ("Distribution email",
         "Envoi SMTP TLS avec piece jointe .docx, corps HTML synthetisant les compteurs "
         "et le top 3 des opportunites. Gestion de groupes To/Cc/Bcc dans l'interface."),
        ("Tableau de bord Gradio",
         "Interface web pour : statut du dernier run, gestion des sources (ajout/edition/test), "
         "validation des cas borderline, apercu/telechargement du rapport, gestion des "
         "destinataires, boutons Run Now et Rebuild Report."),
        ("Observabilite complete",
         "Logs structures JSON via structlog, metriques Prometheus "
         "(notices_found, notices_relevant, crawl_errors, time_to_report), "
         "historique complet des runs en base de donnees, snapshots HTML/PDF conserves 30 jours."),
    ]

    for title, desc in features:
        heading(doc, f">>  {title}", 2)
        normal(doc, desc)

    doc.add_page_break()


# ── section 3 : technologies ──────────────────────────────────────────────────

def build_technologies(doc: Document) -> None:
    heading(doc, "3. Technologies Utilisees")

    categories = [
        ("Langage et Runtime", [
            ("Python 3.11+", "Langage principal ; typage strict (mypy), async natif, pattern matching."),
        ]),
        ("Orchestration IA", [
            ("LangGraph 0.1+",   "Machine a etats pour le pipeline multi-agents (noeuds, transitions, gestion d'erreurs)."),
            ("LangChain 0.2+",   "Abstractions pour les LLM, loaders de documents, outils Fetcher."),
            ("LangChain-Groq",   "Integration Groq API (modele par defaut : llama-3.3-70b-versatile)."),
            ("LangChain-OpenAI", "Integration OpenAI GPT-4o (option alternative)."),
            ("LangChain-Ollama", "Support des modeles locaux via Ollama (mode offline)."),
        ]),
        ("Collecte Web et Documents", [
            ("httpx 0.27+",   "Client HTTP async haute performance avec support de retry/backoff."),
            ("Playwright",    "Rendu JavaScript pour les portails dynamiques (optionnel)."),
            ("selectolax",    "Parser HTML ultra-rapide base sur Modest/Lexbor."),
            ("BeautifulSoup4","Parser HTML de fallback pour les structures complexes."),
            ("pdfminer-six",  "Extraction de texte depuis les PDF natifs."),
            ("Docling 2.60+", "OCR avance (francais) pour les PDF image-only."),
        ]),
        ("NLP et Similarite", [
            ("rapidfuzz 3.5+",        "Calcul de similarite textuelle rapide (deduplication)."),
            ("sentence-transformers", "Encodeurs de phrases pour la recherche semantique (RAG)."),
            ("ChromaDB 1.3+",         "Base de donnees vectorielle pour le pipeline RAG PDF."),
            ("spaCy fr_core_news_md", "Analyse linguistique du francais (optionnel)."),
        ]),
        ("API et Interface Utilisateur", [
            ("FastAPI 0.109+", "API REST backend ; JWT OAuth2 ; documentation OpenAPI auto-generee."),
            ("Uvicorn",        "Serveur ASGI haute performance."),
            ("Gradio 4.44+",   "Interface web admin (dashboard, sources, historique)."),
            ("python-jose",    "Gestion des tokens JWT (authentification securisee)."),
            ("passlib/bcrypt", "Hachage securise des mots de passe."),
        ]),
        ("Donnees et Stockage", [
            ("PostgreSQL 16+",     "Base de donnees relationnelle (sources, runs, notices, fichiers)."),
            ("SQLAlchemy 2.0",     "ORM ; sessions async ; migrations via Alembic."),
            ("Alembic 1.13+",      "Gestion versionnee des migrations de schema."),
            ("MinIO (S3-compat.)", "Stockage objet pour snapshots HTML/PDF et rapports .docx."),
            ("boto3 / aioboto3",   "SDK AWS S3 pour l'interaction avec MinIO."),
        ]),
        ("Rapports et Email", [
            ("python-docx 1.1+", "Generation programmatique de fichiers Word (.docx)."),
            ("SMTP TLS",          "Distribution securisee des rapports ; DKIM/SPF/DMARC recommandes."),
            ("email-validator",   "Validation des adresses email des destinataires."),
        ]),
        ("Configuration et Scheduling", [
            ("Pydantic 2 / pydantic-settings", "Validation de configuration ; chargement .env + settings.yaml."),
            ("PyYAML",           "Parsing du fichier settings.yaml avec substitution de variables d'environnement."),
            ("APScheduler 3.10+","Planificateur cron (par defaut : 07h30 Afrique/Ouagadougou)."),
            ("pytz",             "Gestion des fuseaux horaires."),
            ("Click 8.1+",       "Interface CLI (commandes run-once, test-email, etc.)."),
        ]),
        ("Observabilite", [
            ("structlog 23+",     "Logs structures JSON ; contexte par run."),
            ("prometheus-client", "Exposition de metriques Prometheus (/metrics)."),
        ]),
        ("Qualite et DevOps", [
            ("Docker / Docker Compose", "Conteneurisation multi-services ; healthchecks ; restart policies."),
            ("GitHub Actions",          "CI/CD : lint, tests, build images, deploiement SSH."),
            ("Ruff + Black",            "Linting et formatage automatique du code Python."),
            ("mypy",                    "Verification statique des types (strict)."),
            ("pytest + pytest-cov",     "Tests unitaires et d'integration ; couverture >= 80 % imposee."),
            ("Bandit",                  "Analyse statique de securite (SAST)."),
            ("Poetry",                  "Gestion des dependances et du packaging Python."),
        ]),
    ]

    for cat_title, techs in categories:
        heading(doc, cat_title, 2)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        tbl.columns[0].width = Inches(2.4)
        tbl.columns[1].width = Inches(4.1)

        hdr = tbl.add_row()
        hdr.cells[0].text = "Technologie"
        hdr.cells[1].text = "Role"
        for cell in hdr.cells:
            set_cell_bg(cell, "004C97")
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = BODY_SIZE
            run.font.color.rgb = WHITE

        for i, (tech, role) in enumerate(techs):
            row = tbl.add_row()
            row.cells[0].text = tech
            row.cells[1].text = role
            r = row.cells[0].paragraphs[0].runs[0]
            r.bold = True
            r.font.size = BODY_SIZE
            row.cells[1].paragraphs[0].runs[0].font.size = BODY_SIZE
            if i % 2 == 0:
                set_cell_bg(row.cells[0], "EBF3FB")
                set_cell_bg(row.cells[1], "EBF3FB")

        doc.add_paragraph()

    doc.add_page_break()


# ── section 4 : architecture ──────────────────────────────────────────────────

def build_architecture(doc: Document) -> None:
    heading(doc, "4. Architecture du Systeme")

    normal(doc,
        "TenderAI BF suit une architecture en trois couches separees par des interfaces claires : "
        "la couche presentation (FastAPI + Gradio), la couche metier (pipeline LangGraph) et "
        "la couche donnees (PostgreSQL + MinIO + SMTP)."
    )

    heading(doc, "4.1 Couches Architecturales", 2)

    layers = [
        ("Couche Presentation",
         "FastAPI (port 8000) expose une API REST securisee par JWT avec les routes "
         "/api/v1/runs, /api/v1/sources, /api/v1/reports et /health. "
         "Gradio (port 7860) fournit le tableau de bord d'administration."),
        ("Couche Metier : Pipeline LangGraph",
         "Un StateGraph[TenderAIState] orchestre l'ensemble du traitement. Chaque transition "
         "passe par _route_after_step qui court-circuite vers error_handler si error_occurred=True. "
         "Les erreurs non fatales (ex. : SMTP transitoire apres upload reussi) produisent "
         "le statut completed_with_warnings."),
        ("Couche Donnees",
         "PostgreSQL stocke les metadonnees (sources, runs, notices, fichiers). "
         "MinIO (S3-compatible) conserve les snapshots bruts et les rapports .docx. "
         "Le serveur SMTP assure la distribution des rapports aux destinataires."),
    ]

    for title, desc in layers:
        heading(doc, f">>  {title}", 2)
        normal(doc, desc)

    heading(doc, "4.2 Modele de Donnees", 2)
    normal(doc, "Quatre tables SQLAlchemy principales :")

    tables_info = [
        ("Source",  "Portails a surveiller : URL, selecteurs, rate_limit, enabled, last_seen_at."),
        ("Run",     "Executions du pipeline : timestamps, statut, compteurs JSON, lien vers les logs."),
        ("Notice",  "Appels d'offres individuels : is_relevant, is_duplicate, content_hash, relevance_score, classification_method."),
        ("File",    "Fichiers PDF/DOCX stockes dans MinIO : storage_key, checksum, kind."),
    ]

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Inches(1.5)
    tbl.columns[1].width = Inches(5.0)

    hdr = tbl.add_row()
    hdr.cells[0].text = "Table"
    hdr.cells[1].text = "Description"
    for cell in hdr.cells:
        set_cell_bg(cell, "004C97")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = BODY_SIZE
        run.font.color.rgb = WHITE

    for i, (t, d) in enumerate(tables_info):
        row = tbl.add_row()
        row.cells[0].text = t
        row.cells[1].text = d
        r = row.cells[0].paragraphs[0].runs[0]
        r.bold = True
        r.font.size = BODY_SIZE
        row.cells[1].paragraphs[0].runs[0].font.size = BODY_SIZE
        if i % 2 == 0:
            set_cell_bg(row.cells[0], "EBF3FB")
            set_cell_bg(row.cells[1], "EBF3FB")

    doc.add_paragraph()
    doc.add_page_break()


# ── section 5 : méthodologie ──────────────────────────────────────────────────

def build_methodologie(doc: Document) -> None:
    heading(doc, "5. Methodologie : Pipeline de Traitement")

    normal(doc,
        "Le pipeline LangGraph suit une sequence lineaire de dix noeuds specialises. "
        "Chaque noeud recoit et retourne un objet TenderAIState (modele Pydantic), "
        "garantissant une validation de donnees a chaque etape. "
        "Un mecanisme de routage conditionnel assure la resilience en cas d'erreur."
    )

    heading(doc, "5.1 Noeuds du Pipeline", 2)

    nodes = [
        ("(1) load_sources",
         "Charge depuis PostgreSQL les sources actives avec leur configuration (URL, selecteurs, rate_limit)."),
        ("(2) fetch_listings",
         "Telecharge les pages de listing de chaque source (HTML ou PDF) avec respect du robots.txt, ETag et backoff exponentiel."),
        ("(3) extract_item_links",
         "Applique les selecteurs CSS/XPath/regex pour extraire les URLs individuelles des avis detectes."),
        ("(4) fetch_items",
         "Telecharge le contenu complet de chaque avis (page HTML ou PDF)."),
        ("(5) parse_extract",
         "Parse le contenu brut : selecteurs HTML vers champs structures ; pdfminer-six vers texte ; Docling OCR vers texte pour PDF image. Normalise dates, devises et geographie."),
        ("(6) classify",
         "Evalue la pertinence IT/ingenierie : regles a mots-cles (passe 1) puis LLM si score ambigu (passe 2). Attribue un relevance_score et la methode utilisee."),
        ("(7) deduplicate",
         "Compare chaque notice aux entrees existantes via hash de contenu et/ou similarite textuelle. Fusionne les doublons en conservant toutes les URLs source."),
        ("(8) summarize",
         "Genere des resumes en francais (5 a 8 lignes) via le LLM : objet, entite, budget, eligibilite, pieces cles, echeance, contacts."),
        ("(9) compose_report",
         "Construit le fichier .docx avec python-docx (page de garde, resume executif, fiches AO, annexes) et l'uploade dans MinIO."),
        ("(10) email_report",
         "Envoie le rapport par SMTP TLS avec corps HTML synthetique et piece jointe .docx. Les erreurs SMTP non fatales generent completed_with_warnings."),
    ]

    for name, desc in nodes:
        p = doc.add_paragraph()
        apply_spacing(p)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(3)
        r1 = p.add_run(name + "  ")
        r1.bold = True
        r1.font.color.rgb = PRIMARY_COLOR
        r1.font.size = BODY_SIZE
        rb = p.add_run(desc)
        rb.font.size = BODY_SIZE

    heading(doc, "5.2 Classification Hybride", 2)
    normal(doc, "La classification s'effectue en deux passes complementaires :")
    bullet(doc,
        "Dictionnaire de familles de mots-cles metier "
        "(informatique, logiciel, reseau, fibre, cybersecurite, cloud, data, IA, SIG, "
        "serveur, switch, infogerance, maintenance, hebergement, web/mobile, ERP...). "
        "Rapide et deterministe, couvre 80 a 85 % des cas.",
        "Regles (passe 1) : "
    )
    bullet(doc,
        "Le LLM (llama-3.3-70b via Groq par defaut) traite les notices au score ambigu. "
        "Classification zero-shot avec labels {IT_ENG, NOT_IT} + boost de mots-cles dans "
        "le prompt systeme. Configurable via processing.use_llm_classification.",
        "LLM (passe 2) : "
    )

    heading(doc, "5.3 Deduplication Multi-Strategies", 2)
    normal(doc, "Cinq modes configurables via processing.deduplication_method :")

    dedup = [
        ("hash_only",       "Comparaison du hash SHA-256 du contenu brut. Ultra-rapide."),
        ("similarity_only", "Score rapidfuzz sur titre + entite + date. Tolere les variations redactionnelles."),
        ("hash_similarity", "Combinaison des deux (mode recommande par defaut)."),
        ("llm_only",        "Le LLM juge de l'identite de deux avis. Precis mais couteux."),
        ("hybrid",          "Hash puis similarity puis LLM en cascade. Optimal pour la production."),
    ]
    for mode, desc in dedup:
        bullet(doc, desc, f"{mode} : ")

    heading(doc, "5.4 Planification et Observabilite", 2)
    bullet(doc,
        "APScheduler declenche le pipeline tous les jours a 07h30 (fuseau Afrique/Ouagadougou). "
        "Cron configurable via CRON_SCHEDULE dans settings.yaml.",
        "Scheduling : ")
    bullet(doc,
        "Chaque run cree un enregistrement Run en base (started_at, finished_at, status, counts_json). "
        "Les snapshots HTML/PDF sont conserves 30 jours dans MinIO.",
        "Tracabilite : ")
    bullet(doc,
        "Metriques Prometheus exposees sur /metrics : notices_found, notices_relevant, "
        "emails_sent_ok, crawl_errors, parse_errors, time_to_report.",
        "Metriques : ")
    bullet(doc,
        "Logs JSON structures via structlog avec contexte de run injecte a chaque message.",
        "Logs : ")

    doc.add_page_break()


# ── section 6 : déploiement ───────────────────────────────────────────────────

def build_deploiement(doc: Document) -> None:
    heading(doc, "6. Deploiement et Infrastructure")

    heading(doc, "6.1 Topologie Docker Compose", 2)
    normal(doc, "Cinq services orchestres sur un VPS YULCOM :")

    services = [
        ("api",      "FastAPI + LangGraph pipeline", "8000",     "Orchestrateur principal"),
        ("ui",       "Gradio dashboard",             "7860",     "Interface d'administration"),
        ("worker",   "Taches lourdes PDF/OCR",       "N/A",      "Dechargement CPU intensif"),
        ("postgres", "Base de donnees relationnelle", "5432",    "Metadonnees et historique"),
        ("minio",    "Stockage objet S3-compatible",  "9000/9001","Fichiers et rapports"),
    ]

    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Inches(1.3)
    tbl.columns[1].width = Inches(2.2)
    tbl.columns[2].width = Inches(1.0)
    tbl.columns[3].width = Inches(2.0)

    hdr = tbl.add_row()
    for cell, label in zip(hdr.cells, ["Service", "Role", "Port(s)", "Notes"]):
        cell.text = label
        set_cell_bg(cell, "004C97")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = BODY_SIZE
        run.font.color.rgb = WHITE

    for i, (svc, role, port, note) in enumerate(services):
        row = tbl.add_row()
        row.cells[0].text = svc
        row.cells[1].text = role
        row.cells[2].text = port
        row.cells[3].text = note
        r = row.cells[0].paragraphs[0].runs[0]
        r.bold = True
        r.font.size = BODY_SIZE
        for c in row.cells[1:]:
            c.paragraphs[0].runs[0].font.size = BODY_SIZE
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, "EBF3FB")

    doc.add_paragraph()

    heading(doc, "6.2 CI/CD avec GitHub Actions", 2)
    for step in [
        "Lint (ruff) + verification des types (mypy)",
        "Tests unitaires et d'integration (pytest, couverture >= 80 %)",
        "Analyse de securite SAST (Bandit) + verification des licences",
        "Build et push des images Docker (tags main-<sha>)",
        "Deploiement SSH sur VPS YULCOM",
    ]:
        bullet(doc, step)

    heading(doc, "6.3 Securite", 2)
    for item in [
        "Secrets geres via variables d'environnement Docker (jamais commites)",
        "TENDERAI_JWT_SECRET >= 32 caracteres et TENDERAI_ADMIN_PASSWORD >= 8 caracteres, obligatoires au demarrage",
        "SMTP TLS enforced ; DKIM/SPF/DMARC recommandes",
        "Reverse proxy Nginx avec authentification basique pour l'UI",
        "Rotation des credentials et sauvegardes pg_dump nocturnes",
        "Conformite robots.txt et rate-limiting par domaine",
    ]:
        bullet(doc, item)

    doc.add_page_break()


# ── section 7 : tests & qualité ───────────────────────────────────────────────

def build_tests(doc: Document) -> None:
    heading(doc, "7. Strategie de Tests et Qualite")

    heading(doc, "7.1 Niveaux de Tests", 2)
    tests = [
        ("Unitaires",   "Parsers individuels, normaliseurs de dates/devises, seuils du classifieur, "
                         "logique de deduplication. Isoles, rapides, executes a chaque commit."),
        ("Integration", "Pipeline E2E en mode dry-run avec fixtures enregistrees (VCRpy). "
                         "Tests API FastAPI avec client de test pytest-asyncio."),
        ("Charge",      "Simulation de 200 avis/jour pour verifier un runtime < 15 minutes."),
        ("UAT",         "2 semaines de run en parallele d'un suivi manuel. "
                         "Mesure precision/rappel. Tuning des seuils de classification."),
    ]
    for level, desc in tests:
        bullet(doc, desc, f"{level} : ")

    heading(doc, "7.2 Criteres de Qualite", 2)
    criteria = [
        ("Rappel",          ">= 90 % des AO pertinents detectes sur echantillon 2 semaines"),
        ("Precision",       "Faux positifs <= 10 % apres calibration"),
        ("Performance",     "Pipeline complet <= 15 minutes par run"),
        ("Disponibilite",   "Tolerance aux pannes transitoires avec retry/backoff"),
        ("Couverture code", ">= 80 % (impose par pytest-cov, bloque le CI si non atteint)"),
    ]

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Inches(2.0)
    tbl.columns[1].width = Inches(4.5)

    hdr = tbl.add_row()
    hdr.cells[0].text = "Critere"
    hdr.cells[1].text = "Cible"
    for cell in hdr.cells:
        set_cell_bg(cell, "004C97")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = BODY_SIZE
        run.font.color.rgb = WHITE

    for i, (crit, target) in enumerate(criteria):
        row = tbl.add_row()
        row.cells[0].text = crit
        row.cells[1].text = target
        r = row.cells[0].paragraphs[0].runs[0]
        r.bold = True
        r.font.size = BODY_SIZE
        row.cells[1].paragraphs[0].runs[0].font.size = BODY_SIZE
        if i % 2 == 0:
            set_cell_bg(row.cells[0], "EBF3FB")
            set_cell_bg(row.cells[1], "EBF3FB")

    doc.add_paragraph()
    doc.add_page_break()


# ── section 8 : plan & évolutions ─────────────────────────────────────────────

def build_plan(doc: Document) -> None:
    heading(doc, "8. Plan de Mise en Oeuvre et Evolutions")

    heading(doc, "8.1 Phases de Developpement : Phase 1", 2)
    phases = [
        ("S1 : Fondations",           "Depot Git, Docker Compose, schema PostgreSQL, sources.yaml, crawlers de base, integration MinIO."),
        ("S2 : Parsing et Classifieur","Extraction PDF/HTML, normalisation dates/devises, regles + modele LLM baseline."),
        ("S3 : Resumes et Rapport",    "Prompts FR optimises, template .dotx YULCOM, styles Word, exemples de rapports."),
        ("S4 : Orchestration et Email","Graphe LangGraph complet, retries, SMTP TLS, planification quotidienne APScheduler."),
        ("S5 : Gradio et UAT",         "Interface d'administration, QA borderline, gestion destinataires, tests d'acceptation, go-live."),
    ]
    for sprint, desc in phases:
        bullet(doc, desc, f"{sprint} : ")

    heading(doc, "8.2 Evolutions : Phase 2 (Roadmap)", 2)
    for item in [
        "Scoring et priorisation des AO (adequation profil YULCOM, marge estimee, historique de succes)",
        "Extension multi-pays : UEMOA, CEDEAO (Cote d'Ivoire, Senegal, Mali, Niger...)",
        "Aide a la redaction d'offres (generation de plans, pre-remplissage de dossiers)",
        "Portails avec authentification (Phase 2, identifiants fournis par YULCOM)",
        "Integration CRM pour tracabilite de la reponse aux AO",
        "Application mobile (notifications push lors de nouvelles opportunites)",
    ]:
        bullet(doc, item)

    heading(doc, "8.3 Livrables Finaux", 2)
    for item in [
        "Code source complet (licence interne YULCOM)",
        "Stack Docker Compose + .env.example documente",
        "Interface Gradio deployee sur VPS YULCOM",
        "Modele Word .dotx brande + exemples de rapports",
        "Runbook operationnel + tableau de bord metriques Prometheus",
        "Artefacts de modele et scripts de re-entrainement",
    ]:
        bullet(doc, item)


# ── main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    doc = Document()

    # marges
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    # style de base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = BODY_SIZE
    style.font.color.rgb = DARK_GRAY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = LINE_SPACING

    build_cover(doc)
    add_header(doc)
    add_footer(doc)
    build_presentation(doc)
    build_fonctionnalites(doc)
    build_technologies(doc)
    build_architecture(doc)
    build_methodologie(doc)
    build_deploiement(doc)
    build_tests(doc)
    build_plan(doc)

    out = "TenderAI_BF_Presentation.docx"
    doc.save(out)
    print(f"Document genere : {out}")


if __name__ == "__main__":
    main()
